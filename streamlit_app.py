
import json
import re
import html
import io
from pathlib import Path
from string import Template
from datetime import datetime
import unicodedata as _ud
from typing import Any

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

# --- PATHS (bullet-proof) ---
APP_DIR = Path(__file__).resolve().parent
DICT_DIR = APP_DIR / "dictionaries" / "en_US"   # expects en_US.aff + en_US.dic
RULES_DIR = APP_DIR / "Rules"

# ===========================
# US DICTIONARY (PURE PYTHON)
# ===========================
try:
    from spylls.hunspell import Dictionary
except Exception:
    Dictionary = None  # we'll handle gracefully below

# Toggle loader debug messages (False for production)
DEBUG_DICT_LOAD = False

# Spell-check tuning
SPELL_MIN_LEN = 3
SPELL_WHITELIST = {
    "Textile", "Exchange", "Textile Exchange", "degrowth",
}

# --- DEFAULT ALL-CAPS whitelist (acronyms you don't want flagged) ---
DEFAULT_CAPS_WHITELIST = {
    "US", "USA", "UK", "EU", "UN", "UNEP", "UNDP", "UNFCCC", "IPCC",
    "OECD", "NASA", "NATO", "ISO", "SDG", "SDGs", "ESG", "CSR", "FAQ",
    "PDF", "CSV", "DOCX", "XLSX", "HTTP", "HTTPS", "API", "CO", "GHG",
}

# --- Google Sheets deps (Pylance-friendly fallbacks) ---
gspread: Any = None
Credentials: Any = None
WorksheetNotFound: type[Exception] = Exception
try:
    import gspread as _gspread
    from google.oauth2.service_account import Credentials as _Credentials
    from gspread.exceptions import WorksheetNotFound as _WorksheetNotFound

    gspread = _gspread
    Credentials = _Credentials
    WorksheetNotFound = _WorksheetNotFound
except Exception:
    pass

# ===========================
# PAGE CONFIG
# ===========================
st.set_page_config(page_title="Textile Exchange Style Guide Buddy", layout="wide")
st.title("📘 Textile Exchange Style Guide Buddy")

# ===========================
# US DICTIONARY LOADER (cache-safe)
# ===========================
def _dict_cache_key() -> tuple:
    """Cache key from existence/size/mtime of the dict files."""
    aff = DICT_DIR / "en_US.aff"
    dic = DICT_DIR / "en_US.dic"

    def _meta(p: Path):
        try:
            return (
                p.exists(),
                p.stat().st_size if p.exists() else 0,
                int(p.stat().st_mtime) if p.exists() else 0,
            )
        except Exception:
            return (False, 0, 0)

    return _meta(aff) + _meta(dic)

@st.cache_resource(show_spinner=False)
def get_us_dictionary(cache_key: tuple):
    """Load Hunspell en_US via spylls. Returns Dictionary or None."""
    if Dictionary is None:
        if DEBUG_DICT_LOAD:
            st.error("spylls import failed: `from spylls.hunspell import Dictionary` returned None.")
        return None

    try:
        aff = DICT_DIR / "en_US.aff"
        dic = DICT_DIR / "en_US.dic"

        if DEBUG_DICT_LOAD:
            st.info(f"DEBUG: aff={aff} (exists={aff.exists()}, size={aff.stat().st_size if aff.exists() else '—'})")
            st.info(f"DEBUG: dic={dic} (exists={dic.exists()}, size={dic.stat().st_size if dic.exists() else '—'})")
            st.info(f"DEBUG: cache_key={cache_key}")

        if not (aff.exists() and dic.exists()):
            if DEBUG_DICT_LOAD:
                st.error("Dictionary files not found at expected locations.")
            return None

        # Pass a single base path (no extension)
        d = Dictionary.from_files(str(DICT_DIR / "en_US"))
        if DEBUG_DICT_LOAD:
            st.success("DEBUG: spylls Dictionary loaded successfully.")
        return d

    except Exception as e:
        if DEBUG_DICT_LOAD:
            st.exception(e)
        return None

US_DICT = get_us_dictionary(_dict_cache_key())

# (Status banner)
SHOW_SPELLCHECK_STATUS = True
if SHOW_SPELLCHECK_STATUS:
    status = "ACTIVE" if US_DICT else "INACTIVE (missing dictionaries)"
    st.caption(f"Spellcheck (spylls en_US): {status}")

# Reload button
if st.button("♻ Reload US dictionary", type="secondary"):
    try:
        get_us_dictionary.clear()
        st.toast("US dictionary cache cleared. Rebuilding…")
    except Exception:
        pass

if SHOW_SPELLCHECK_STATUS and not US_DICT:
    st.warning(
        "US dictionary files not found at `dictionaries/en_US/`. "
        "Spellcheck is disabled until `en_US.aff` and `en_US.dic` are present."
    )

# ===========================
# RULES STORAGE (SHEETS or LOCAL)
# ===========================
RULES_FILE = RULES_DIR / "Textile_Exchange_Style_Guide_STRICT.json"
CAPS_FILE  = RULES_DIR / "caps_whitelist.json"  # local fallback for acronym whitelist

def _has_service_account_in_secrets() -> bool:
    try:
        if "google_service_account_json" in st.secrets:
            return True
        if "google_service_account" in st.secrets:
            return True
        root = st.secrets
        return (
            isinstance(root.get("type"), str)
            and root.get("type") == "service_account"
            and isinstance(root.get("client_email"), str)
            and isinstance(root.get("private_key"), str)
        )
    except Exception:
        return False

SHEETS_ENABLED = (
    (gspread is not None)
    and _has_service_account_in_secrets()
    and "gsheets" in st.secrets
    and "SPREADSHEET_ID" in st.secrets["gsheets"]
)
CAPS_WS_NAME = st.secrets["gsheets"].get("CAPS_WS_NAME", "caps_whitelist") if SHEETS_ENABLED else "caps_whitelist"

def _coerce_bool(v) -> bool:
    if isinstance(v, bool): return v
    if isinstance(v, (int, float)): return v != 0
    if isinstance(v, str): return v.strip().lower() in ("true", "t", "yes", "y", "1", "on")
    return False

# ---- Local rules ----
def load_rules_local() -> dict:
    if RULES_FILE.exists():
        data = json.loads(RULES_FILE.read_text(encoding="utf-8"))
        data.setdefault("style_guide_rule", [])
        data.setdefault("style_guide_caution", [])
        for cat in ("style_guide_rule", "style_guide_caution"):
            for it in data.get(cat, []):
                it["case_sensitive"] = _coerce_bool(it.get("case_sensitive"))
        return data
    return {"style_guide_rule": [], "style_guide_caution": []}

def save_rules_local(rules: dict):
    RULES_DIR.mkdir(exist_ok=True, parents=True)
    RULES_FILE.write_text(json.dumps(rules, indent=2, ensure_ascii=False), encoding="utf-8")

# ---- Local caps whitelist ----
def load_caps_whitelist_local() -> set[str]:
    try:
        if CAPS_FILE.exists():
            items = json.loads(CAPS_FILE.read_text(encoding="utf-8"))
            return {str(x).strip().upper() for x in items if str(x).strip()}
    except Exception:
        pass
    return set()

def save_caps_whitelist_local(acros: set[str]):
    RULES_DIR.mkdir(exist_ok=True, parents=True)
    CAPS_FILE.write_text(json.dumps(sorted(acros), indent=2, ensure_ascii=False), encoding="utf-8")

# ---- PEM canonicalizer (kept) ----
def _canonicalize_pem(pem: str) -> str:
    """
    Rebuild a well-formed PKCS#8 PEM from whatever was pasted:
    - Keep BEGIN/END lines
    - Strip non-base64 junk
    - Re-wrap base64 to 64-char lines
    - Fix padding if the length mod 4 is off
    """
    if not pem or "BEGIN PRIVATE KEY" not in pem or "END PRIVATE KEY" not in pem:
        return pem

    import re as _re, textwrap
    pem = pem.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.strip() for ln in pem.splitlines()]

    header = "-----BEGIN PRIVATE KEY-----"
    footer = "-----END PRIVATE KEY-----"

    # Pull out base64 body lines (everything not header/footer and not empty)
    body = "".join(ln for ln in lines if ln and ln != header and ln != footer)
    # Remove anything that's not base64 alphabet or padding
    body = _re.sub(r"[^A-Za-z0-9+/=]", "", body)

    # Fix padding if required
    rem = len(body) % 4
    if rem:
        body += "=" * (4 - rem)

    wrapped = "\n".join(textwrap.wrap(body, 64))
    return f"{header}\n{wrapped}\n{footer}\n"

# ---- Google Sheets: auth helpers ----
GS_EXPECTED_COLS = ["category", "match", "replace_with", "message", "case_sensitive", "updated_at"]

def _sa_info_from_secrets() -> dict:
    if "google_service_account_json" in st.secrets:
        raw = st.secrets["google_service_account_json"]
        sa_info = json.loads(raw)
    elif "google_service_account" in st.secrets:
        sa_info = dict(st.secrets["google_service_account"])
    else:
        sa_info = dict(st.secrets)  # raw at root

    pk = sa_info.get("private_key", "")
    if "BEGIN PRIVATE KEY" in pk:
        # normalize newline escapes -> real newlines
        sa_info["private_key"] = pk.replace("\\n", "\n").replace("\r\n", "\n").replace("\r", "\n")
        # canonicalize
        sa_info["private_key"] = _canonicalize_pem(sa_info["private_key"])
    return sa_info


def _validate_private_key_pem(pem: str):
    """
    Validate the PEM just enough to catch obvious paste errors, but be tolerant
    of harmless whitespace artifacts introduced by Secrets UI.
    """
    if not pem or "BEGIN PRIVATE KEY" not in pem or "END PRIVATE KEY" not in pem:
        raise ValueError("Private key is missing BEGIN/END PRIVATE KEY markers.")

    # Extract the base64 body (lines between the markers)
    lines = [ln.rstrip("\r\n") for ln in pem.splitlines()]
    # Keep everything that's not a header/footer
    body_lines = [ln.strip() for ln in lines if not ln.startswith("-----") and ln.strip()]
    body = "".join(body_lines)

    # Remove any characters that are not in base64 alphabet (defensive)
    import re as _re, base64, binascii
    cleaned = _re.sub(r"[^A-Za-z0-9+/=]", "", body)

    # Optional sanity: strip trailing padding noise beyond two '='
    cleaned = _re.sub(r"=+$", lambda m: "=" * min(2, len(m.group(0))), cleaned)

    try:
        # Be permissive here; we already sanitized non-base64 chars.
        base64.b64decode(cleaned, validate=False)
    except binascii.Error as e:
        raise ValueError(f"Private key Base64 decoding failed (likely newline/quoting issue): {e}") from e


@st.cache_resource(show_spinner=False)
def get_gspread_client():
    if gspread is None or Credentials is None:
        raise RuntimeError("Google Sheets dependencies are not installed in this environment.")
    scopes = ["https://www.googleapis.com/auth/spreadsheets"]
    sa_info = _sa_info_from_secrets()
    _validate_private_key_pem(sa_info.get("private_key", ""))
    creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
    return gspread.authorize(creds)

@st.cache_resource(show_spinner=False)
def get_or_create_worksheet():
    if gspread is None:
        raise RuntimeError("gspread is not available.")
    gc = get_gspread_client()
    sh = gc.open_by_key(st.secrets["gsheets"]["SPREADSHEET_ID"])
    ws_name = st.secrets["gsheets"].get("WORKSHEET_NAME", "rules")
    try:
        ws = sh.worksheet(ws_name)
    except WorksheetNotFound:
        ws = sh.add_worksheet(title=ws_name, rows=100, cols=10)
        ws.update("A1", [GS_EXPECTED_COLS])
    return ws

@st.cache_data(ttl=60, show_spinner=False)
def load_rules_sheets() -> dict:
    ws = get_or_create_worksheet()
    records = ws.get_all_records()
    out = {"style_guide_rule": [], "style_guide_caution": []}
    for r in records:
        cat = (r.get("category") or "").strip()
        if cat not in ("style_guide_rule", "style_guide_caution"):
            continue
        item = {
            "match": r.get("match") or "",
            "replace_with": r.get("replace_with") or None,
            "message": r.get("message") or "",
            "case_sensitive": _coerce_bool(r.get("case_sensitive")),
        }
        out[cat].append(item)
    return out

def save_rules_sheets(rules: dict):
    ws = get_or_create_worksheet()
    rows = []
    now = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    for cat in ("style_guide_rule", "style_guide_caution"):
        for r in rules.get(cat, []):
            rows.append([
                cat,
                r.get("match") or "",
                r.get("replace_with") or "",
                r.get("message") or "",
                bool(r.get("case_sensitive", False)),
                now,
            ])
    ws.clear()
    ws.update("A1", [GS_EXPECTED_COLS] + rows)
    try:
        load_rules_sheets.clear()
    except Exception:
        pass

# ---- Google Sheets: caps whitelist ----
@st.cache_resource(show_spinner=False)
def get_or_create_caps_ws():
    gc = get_gspread_client()
    sh = gc.open_by_key(st.secrets["gsheets"]["SPREADSHEET_ID"])
    try:
        ws = sh.worksheet(CAPS_WS_NAME)
    except WorksheetNotFound:
        ws = sh.add_worksheet(title=CAPS_WS_NAME, rows=100, cols=2)
        ws.update("A1", [["acronym", "updated_at"]])
    return ws

@st.cache_data(ttl=60, show_spinner=False)
def load_caps_whitelist_sheets() -> set[str]:
    ws = get_or_create_caps_ws()
    records = ws.get_all_records()
    return {(r.get("acronym") or "").strip().upper() for r in records if (r.get("acronym") or "").strip()}

def save_caps_whitelist_sheets(acros: set[str]):
    ws = get_or_create_caps_ws()
    now = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    rows = [["acronym", "updated_at"]] + [[a, now] for a in sorted(acros)]
    ws.clear()
    ws.update("A1", rows)
    try:
        load_caps_whitelist_sheets.clear()
    except Exception:
        pass

# ---- Choosers (Rules & Caps whitelist) ----
def load_rules() -> dict:
    if SHEETS_ENABLED:
        try:
            return load_rules_sheets()
        except Exception:
            st.warning("Google Sheets unavailable or misconfigured. Falling back to local JSON.")
            return load_rules_local()
    return load_rules_local()

def save_rules(rules: dict):
    if SHEETS_ENABLED:
        try:
            save_rules_sheets(rules)
            return
        except Exception:
            st.warning("Could not save to Google Sheets. Saved to local JSON instead.")
    save_rules_local(rules)

def load_caps_whitelist() -> set[str]:
    if SHEETS_ENABLED:
        try:
            return load_caps_whitelist_sheets()
        except Exception:
            st.warning("Google Sheets unavailable; using local caps whitelist.")
            return load_caps_whitelist_local()
    return load_caps_whitelist_local()

def save_caps_whitelist(acros: set[str]):
    if SHEETS_ENABLED:
        try:
            save_caps_whitelist_sheets(acros)
            return
        except Exception:
            st.warning("Could not save caps whitelist to Google Sheets. Saved locally instead.")
    save_caps_whitelist_local(acros)

# ===========================
# SESSION STATE
# ===========================
def ensure_state():
    if "rules" not in st.session_state:
        st.session_state.rules = load_rules()
    if "edit_rule" not in st.session_state:
        st.session_state.edit_rule = None
    if "caps_whitelist" not in st.session_state:
        st.session_state.caps_whitelist = load_caps_whitelist()

ensure_state()

# ===========================
# MATCHING & SPELL-CHECK HELPERS
# ===========================

def _is_block_mostly_all_caps(text: str) -> bool:
    """
    Returns True if this block (paragraph/cell) is essentially ALL‑CAPS,
    even if sentence splitting fails (common with .docx run artifacts).
    Heuristic:
      - long blocks: >= 10 letters and >= 80% uppercase
      - short headings: >= 6 letters and 100% uppercase
    """
    if not text:
        return False

    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False

    upper = sum(1 for ch in letters if ch.isupper())
    n = len(letters)

    long_mostly_caps = (n >= 10 and upper / n >= 0.80)
    short_all_caps   = (n >= 6 and upper == n)

    return long_mostly_caps or short_all_caps


def _suppress_cap_words_inside_cap_sent(caps_sent_issues: list[dict], caps_word_issues: list[dict]) -> list[dict]:
    """
    If a sentence/line is ALL‑CAPS (yellow caution), suppress individual ALL‑CAPS
    word cautions fully contained within any such sentence/line span.
    """
    if not caps_sent_issues or not caps_word_issues:
        return caps_word_issues

    # Build list of sentence spans
    sent_spans = [(int(i["start"]), int(i["end"])) for i in caps_sent_issues]

    def _inside_any_sentence(w_start: int, w_end: int) -> bool:
        for s_start, s_end in sent_spans:
            # fully contained
            if w_start >= s_start and w_end <= s_end:
                return True
        return False

    filtered = []
    for w in caps_word_issues:
        w_start, w_end = int(w.get("start", -1)), int(w.get("end", -1))
        if not _inside_any_sentence(w_start, w_end):
            filtered.append(w)
    return filtered


def _normalize_for_match(s: str) -> str:
    if not s:
        return ""
    s = html.unescape(s)
    s = _ud.normalize("NFC", s)
    return s

def _needs_word_boundaries(token: str) -> bool:
    return re.fullmatch(r"[A-Za-z0-9]+(?:['’-][A-Za-z0-9]+)*", token) is not None

def find_matches(text, rules, location, prefix):
    matches = []
    for rule in rules:
        word = rule.get("match")
        if not word:
            continue
        flags = 0 if rule.get("case_sensitive") else re.IGNORECASE

        use_regex = isinstance(word, str) and len(word) >= 2 and word.startswith("/") and word.endswith("/")
        if use_regex:
            pattern = word[1:-1]
        else:
            if _needs_word_boundaries(word):
                pattern = rf"\b{re.escape(word)}\b"
            else:
                pattern = re.escape(word)

        for m in re.finditer(pattern, text, flags):
            matches.append({
                "start": m.start(),
                "end": m.end(),
                "issue": m.group(),
                "replacement": rule.get("replace_with"),
                "explanation": rule.get("message"),
                "category": rule["category"],
                "location": location,
            })

    matches.sort(key=lambda x: (x["start"], x["end"]))
    for i, m in enumerate(matches, 1):
        m["anchor"] = f"{prefix}_m{i}"
    return matches

# ===== UK→US ENFORCEMENT =====
UK_US_MAP = {
    "organisation": "organization", "organisations": "organizations",
    "organisational": "organizational",
    "colour": "color", "colours": "colors",
    "flavour": "flavor", "flavours": "flavors",
    "behaviour": "behavior", "behaviours": "behaviors",
    "favourite": "favorite", "favourites": "favorites",
    "neighbour": "neighbor", "neighbours": "neighbors",
    "neighbourhood": "neighborhood", "neighbourhoods": "neighborhoods",
    "programme": "program", "programmes": "programs",
    "catalogue": "catalog", "catalogues": "catalogs",
    "dialogue": "dialog", "dialogues": "dialogs",
    "traveller": "traveler", "travellers": "travelers",
    "modelling": "modeling", "modelled": "modeled",
    "labour": "labor", "labourers": "laborers",
    "centre": "center", "centres": "centers",
    "theatre": "theater", "theatres": "theaters",
    "metre": "meter", "metres": "meters",
    "litre": "liter", "litres": "liters",
    "kilometre": "kilometer", "kilometres": "kilometers",
    "tonne": "ton",
    "defence": "defense",
    "licence": "license", "licences": "licenses",
    "offence": "offense",
    "practise": "practice", "practised": "practiced", "practising": "practicing",
    "grey": "gray",
    "aluminium": "aluminum",
    "cheque": "check", "cheques": "checks",
    "aeroplane": "airplane",
    "tyre": "tire", "tyres": "tires",
    "rumour": "rumor", "humour": "humor", "vapour": "vapor", "odour": "odor",
    "mould": "mold", "moulding": "molding",
    "jewellery": "jewelry",
    "storey": "story", "storeys": "stories",
    "cosy": "cozy",
}
ISE_EXCEPTIONS = {
    "advise", "arise", "chastise", "comprise", "compromise", "demise", "devise",
    "disguise", "enterprise", "exercise", "franchise", "improvise", "merchandise",
    "paradise", "premise", "precise", "revise", "rise", "surprise", "sunrise",
    "promise", "advertise",
}

def _preserve_case(us_replacement: str, original: str) -> str:
    if original.isupper():
        return us_replacement.upper()
    if original[:1].isupper():
        return us_replacement[:1].upper() + us_replacement[1:]
    return us_replacement

def _uk_to_us_if_applicable(token: str) -> str | None:
    base = token.lower()
    if base in UK_US_MAP:
        return _preserve_case(UK_US_MAP[base], token)
    if base.endswith("our") and len(base) > 3:
        return _preserve_case(base[:-3] + "or", token)
    if base.endswith("re") and len(base) > 2:
        if re.search(r"(cent|theat|met|litr|kilometr|metr|kilom|centim)re$", base):
            return _preserve_case(base[:-2] + "er", token)
    if re.search(r"(?:ise|ised|ising)$", base) and base not in ISE_EXCEPTIONS:
        return _preserve_case(
            re.sub(r"ised$", "ized", re.sub(r"ising$", "izing", re.sub(r"ise$", "ize", base))),
            token
        )
    if re.search(r"(?:alyse|nalyse|paralyse|catalyse|dialyse|electrolyse)$", base):
        return _preserve_case(re.sub(r"lyse$", "lyze", base), token).replace("alyse", "alyze")
    return None

def find_non_us_words(text: str, location: str, prefix: str):
    issues = []
    token_re = re.compile(r"[A-Za-z][A-Za-z’'-]*")
    for i, m in enumerate(token_re.finditer(text), 1):
        tok = m.group()
        repl = _uk_to_us_if_applicable(tok)
        if repl is not None and repl != tok:
            issues.append({
                "start": m.start(), "end": m.end(), "issue": tok, "replacement": repl,
                "explanation": "Use US English spelling.", "category": "style_guide_rule",
                "location": location, "anchor": f"{prefix}_us{i}",
            })
    return issues

def find_non_dictionary_words(text: str, location: str, prefix: str):
    issues = []
    if not US_DICT:
        return issues
    token_re = re.compile(r"[A-Za-z][A-Za-z’'-]*")
    i = 0
    for m in token_re.finditer(text):
        token = m.group()
        base = token.strip()
        if len(base) < SPELL_MIN_LEN: continue
        if any(ch.isdigit() for ch in base): continue
        if base in SPELL_WHITELIST: continue
        def _ok(w: str) -> bool:
            try:
                return bool(US_DICT.lookup(w))
            except Exception:
                return False
        if not (_ok(base) or _ok(base.lower()) or _ok(base.capitalize())):
            i += 1
            try:
                suggestions = US_DICT.suggest(base)[:3]
            except Exception:
                suggestions = []
            expl = "Word not in US English dictionary"
            if suggestions:
                expl += f" (e.g., {', '.join(suggestions)})"
            issues.append({
                "start": m.start(), "end": m.end(), "issue": token, "replacement": None,
                "explanation": expl, "category": "style_guide_rule",
                "location": location, "anchor": f"{prefix}_dict{i}",
            })
    return issues

# ----- ALL-CAPS SENTENCE CAUTION -----

def _iter_sentences_with_spans(text: str):
    """
    Yield (start, end, sentence_text) using a tolerant splitter:
    - Primary split on [.?!] followed by whitespace.
    - Also split on explicit newlines so ALL‑CAPS lines without punctuation are considered.
    - Always return the trailing chunk, even without terminal punctuation.
    """
    if not text:
        return [(0, 0, "")]

    spans = []
    last = 0

    # First pass: split on punctuation + whitespace (e.g., ". " or "!  ")
    for m in re.finditer(r"([.!?])\s+", text):
        end = m.end()
        spans.append((last, end, text[last:end]))
        last = end

    # Handle the tail (text after the last punctuation split)
    if last < len(text):
        tail = text[last:]

        # Second pass within the tail: split on explicit newlines
        line_start_abs = last
        for ln in re.finditer(r"\n+", tail):
            line_end_abs = last + ln.start()
            if line_end_abs > line_start_abs:
                spans.append((line_start_abs, line_end_abs, text[line_start_abs:line_end_abs]))
            line_start_abs = last + ln.end()

        # Final trailing line/chunk
        if line_start_abs < len(text):
            spans.append((line_start_abs, len(text), text[line_start_abs:]))

    # Fallback: entire text if nothing matched
    if not spans:
        spans = [(0, len(text), text)]
    return spans



def find_all_caps_sentences(text: str, location: str, prefix: str):
    """
    Detect contiguous ALL‑CAPS spans (regardless of .docx run boundaries)
    and return ONE caution per span.
    """
    issues = []
    token_re = re.compile(r"\b[A-Za-z][A-Za-z’'-]*\b")
    
    tokens = [(m.start(), m.end(), m.group()) for m in token_re.finditer(text)]
    if not tokens:
        return issues

    

def find_all_caps_sentences(text: str, location: str, prefix: str):
    """
    Detect contiguous ALL‑CAPS spans (regardless of .docx run boundaries)
    and return ONE caution per span.
    """
    # Normalise hidden breaks
    text = text.replace("\r", "").replace("\n", " ")

    issues = []
    token_re = re.compile(r"\b[A-Za-z][A-Za-z’'-]*\b")

    tokens = [(m.start(), m.end(), m.group()) for m in token_re.finditer(text)]
    if not tokens:
        return []

    # Whitelist-aware caps detector
    def is_caps_word(tok: str):
        wl = DEFAULT_CAPS_WHITELIST | st.session_state.get("caps_whitelist", set())
        letters = "".join(ch for ch in tok if ch.isalpha())
        if len(letters) < 2:
            return False
        if letters in wl:
            return False
        return letters.isupper()

    # Build spans of consecutive ALL‑CAPS tokens
    spans = []
    span_start = None
    for i, (s, e, tok) in enumerate(tokens):
        if is_caps_word(tok):
            if span_start is None:
                span_start = i
        else:
            if span_start is not None:
                spans.append((span_start, i - 1))
                span_start = None
    if span_start is not None:
        spans.append((span_start, len(tokens) - 1))

    final_spans = []
    for si, ei in spans:
        span_tokens = tokens[si:ei + 1]

        letters = [ch for _, _, tok in span_tokens for ch in tok if ch.isalpha()]
        if len(letters) < 6:
            continue

        start = span_tokens[0][0]
        end   = span_tokens[-1][1]

        # Tighten boundaries
        while start < end and not text[start].isalpha():
            start += 1
        while end > start and not text[end-1].isalpha():
            end -= 1

        final_spans.append((start, end))

    # Build issues
    out = []
    for idx, (s, e) in enumerate(final_spans, 1):
        chunk = text[s:e]
        out.append({
            "start": s,
            "end": e,
            "issue": chunk[:80] + ("…" if len(chunk) > 80 else ""),
            "replacement": None,
            "explanation": "Avoid ALL‑CAPS sentences/lines; use normal capitalization.",
            "category": "style_guide_caution",
            "location": location,
            "anchor": f"{prefix}_caps{idx}",
        })

    return out

    # Walk tokens & form spans of consecutive ALL‑CAPS tokens
    spans = []
    span_start = None
    for i, (s, e, tok) in enumerate(tokens):
        if is_caps_word(tok):
            if span_start is None:
                span_start = i
        else:
            if span_start is not None:
                spans.append((span_start, i - 1))
                span_start = None
    if span_start is not None:
        spans.append((span_start, len(tokens) - 1))

    # Convert spans → (start_idx, end_idx in the original text)
    final_spans = []
    for si, ei in spans:
        span_tokens = tokens[si:ei + 1]
        letters = [ch for _, _, tok in span_tokens for ch in tok if ch.isalpha()]
        if len(letters) < 6:
            continue   # ignore small sequences like “THE”
        # build final span
        start = span_tokens[0][0]
        end   = span_tokens[-1][1]
        final_spans.append((start, end))

    # Create issues
    for idx, (s, e) in enumerate(final_spans, 1):
        chunk = text[s:e]
        issues.append({
            "start": s,
            "end": e,
            "issue": chunk[:80] + ("…" if len(chunk) > 80 else ""),
            "replacement": None,
            "explanation": "Avoid ALL‑CAPS sentences/lines; use normal capitalization.",
            "category": "style_guide_caution",
            "location": location,
            "anchor": f"{prefix}_caps{idx}",
        })

    return issues

    # --------------------------------------

    # Normal per‑sentence logic
    for idx, (s, e, sent) in enumerate(_iter_sentences_with_spans(text), 1):
        chunk = sent.strip()
        if not chunk:
            continue

        letters = [ch for ch in chunk if ch.isalpha()]
        if not letters:
            continue

        upper = sum(1 for ch in letters if ch.isupper())
        n = len(letters)

        long_mostly_caps = (n >= 10 and upper / n >= 0.80)
        short_all_caps   = (n >= 6  and upper == n)

        if long_mostly_caps or short_all_caps:
            issues.append({
                "start": s,
                "end": e,
                "issue": chunk[:80] + ("…" if len(chunk) > 80 else ""),
                "replacement": None,
                "explanation": "Avoid ALL‑CAPS sentences/lines; use normal capitalization.",
                "category": "style_guide_caution",
                "location": location,
                "anchor": f"{prefix}_caps{idx}",
            })

    return issues



# ----- ALL-CAPS WORD CAUTION -----
def find_all_caps_words(text: str, location: str, prefix: str):
    issues = []
    token_re = re.compile(r"\b[A-Za-z][A-Za-z’'-]*\b")
    i = 0
    wl = DEFAULT_CAPS_WHITELIST | st.session_state.get("caps_whitelist", set())
    for m in token_re.finditer(text):
        tok = m.group()
        if any(ch.isdigit() for ch in tok):
            continue
        letters_only = "".join(ch for ch in tok if ch.isalpha())
        if len(letters_only) < 2:
            continue
        if letters_only in wl:
            continue
        if letters_only.isupper():
            i += 1
            issues.append({
                "start": m.start(), "end": m.end(), "issue": tok, "replacement": None,
                "explanation": "Avoid ALL‑CAPS words (except acronyms); use normal capitalization.",
                "category": "style_guide_caution", "location": location, "anchor": f"{prefix}_capw{i}",
            })
    return issues

# ----- DEDUPE & SORT HELPERS -----
def _severity_key(item: dict):
    sev = 0 if item.get("category") == "style_guide_rule" else 1
    return (sev, str(item.get("location", "")), int(item.get("start", 0)), int(item.get("end", 0)))

def _dedupe_by_span(issues: list[dict]) -> list[dict]:
    def _pri(a: str) -> int:
        if "_us" in a:   return 0
        if "_m" in a:    return 1
        if "_dict" in a: return 2
        if "_caps" in a or "_capw" in a: return 3
        return 9
    sorted_issues = sorted(
        issues,
        key=lambda i: (_pri(i.get("anchor", "")), int(i.get("start", 0)), int(i.get("end", 0)))
    )
    kept, seen_spans = [], set()
    for it in sorted_issues:
        span = (int(it.get("start", -1)), int(it.get("end", -1)))
        if span not in seen_spans:
            seen_spans.add(span)
            kept.append(it)
    kept = sorted(kept, key=_severity_key)
    return kept

# ===========================
# INLINE CHECKER RENDERING (.docx)
# ===========================
HIGHLIGHT_STYLE = {
    "style_guide_rule": "border-bottom:2px solid #ff4d4d;",
    "style_guide_caution": "border-bottom:2px solid #ffcc00;",
}

def flatten_rules():
    rules_state = st.session_state.get("rules", {"style_guide_rule": [], "style_guide_caution": []})
    out = []
    for cat in ("style_guide_rule", "style_guide_caution"):
        for r in rules_state.get(cat, []):
            if r.get("match"):
                out.append({**r, "category": cat})
    return out



def render_text(text, matches):
    """
    Safe renderer: no duplication, no broken spans, no misaligned indexes.
    This version advances the cursor correctly and never re-appends matched text.
    """
    if not matches:
        return html.escape(text)

    # Sort matches in case they're not already
    matches = sorted(matches, key=lambda m: m["start"])

    out = []
    cursor = 0

    for m in matches:
        start = m["start"]
        end   = m["end"]

        if start < cursor:
            # overlapping or nested match – skip
            continue

        # Append text before the highlight
        out.append(html.escape(text[cursor:start]))

        tooltip = html.escape(
            f"Issue: {m['issue']}\n"
            f"Replacement: {m['replacement'] or '—'}\n"
            f"{m['explanation']}"
        )

        # Append highlighted section
        out.append(
            f'<span id="{m["anchor"]}" '
            f'style="{HIGHLIGHT_STYLE[m["category"]]}" '
            f'title="{tooltip}">'
            f'{html.escape(text[start:end])}'
            f'</span>'
        )

        # 🔑 Critical fix: ADVANCE CURSOR PAST THE MATCH
        cursor = end

    # Append remaining text AFTER last match
    if cursor < len(text):
        out.append(html.escape(text[cursor:]))

    return "".join(out)


def iter_blocks(doc):
    p_i = t_i = 0
    for el in doc.element.body.iterchildren():
        if isinstance(el, CT_P):
            yield "p", doc.paragraphs[p_i]
            p_i += 1
        elif isinstance(el, CT_Tbl):
            yield "tbl", doc.tables[t_i]
            t_i += 1


def analyze_inline(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    rules = flatten_rules()
    left_parts, issues = [], []

    para_i = tbl_i = 0
    for kind, block in iter_blocks(doc):
        if kind == "p":
            para_i += 1
            text = _normalize_for_match(block.text or "")
            loc = f"Paragraph {para_i}"

            rule_matches = find_matches(text, rules, loc, f"p{para_i}")
            dict_issues  = find_non_dictionary_words(text, loc, f"p{para_i}")
            ukus_issues  = find_non_us_words(text, loc, f"p{para_i}")
            caps_sent    = find_all_caps_sentences(text, loc, f"p{para_i}")
            caps_words   = find_all_caps_words(text, loc, f"p{para_i}")

            # Suppress word-level ALL‑CAPS cautions if a sentence/line is flagged,
            # or if the whole paragraph looks ALL‑CAPS (belt-and-braces for .docx artifacts)
            caps_words = _suppress_cap_words_inside_cap_sent(caps_sent, caps_words)
            if caps_sent or _is_block_mostly_all_caps(text):
                caps_words = []

            combined = _dedupe_by_span(rule_matches + ukus_issues + dict_issues + caps_sent + caps_words)
            issues.extend(combined)
            left_parts.append(f"<p>{render_text(text, combined) or '&nbsp;'}</p>")

        else:
            tbl_i += 1
            rows_html = []
            for r, row in enumerate(block.rows, 1):
                cells = []
                for c, cell in enumerate(row.cells, 1):
                    text = _normalize_for_match(cell.text or "")
                    loc = f"Table {tbl_i}, row {r}, col {c}"

                    rule_matches = find_matches(text, rules, loc, f"t{tbl_i}_{r}_{c}")
                    dict_issues  = find_non_dictionary_words(text, loc, f"t{tbl_i}_{r}_{c}")
                    ukus_issues  = find_non_us_words(text, loc, f"t{tbl_i}_{r}_{c}")
                    caps_sent    = find_all_caps_sentences(text, loc, f"t{tbl_i}_{r}_{c}")
                    caps_words   = find_all_caps_words(text, loc, f"t{tbl_i}_{r}_{c}")

                    caps_words = _suppress_cap_words_inside_cap_sent(caps_sent, caps_words)
                    if caps_sent or _is_block_mostly_all_caps(text):
                        caps_words = []

                    combined = _dedupe_by_span(rule_matches + ukus_issues + dict_issues + caps_sent + caps_words)
                    issues.extend(combined)
                    cells.append(f"<td>{render_text(text, combined) or '&nbsp;'}</td>")

                rows_html.append(f"<tr>{''.join(cells)}</tr>")
            left_parts.append(f"<table class='doc-table'><tbody>{''.join(rows_html)}</tbody></table>")

    left_html = "".join(left_parts) or "<p>&nbsp;</p>"

    # Right-hand list (sorted)
    issues_sorted = sorted(issues, key=_severity_key)
    right_items = []
    for i in issues_sorted:
        color = "#ff4d4d" if i["category"] == "style_guide_rule" else "#ffcc00"
        right_items.append(
            f"<a href=\"#{i['anchor']}\" class=\"issue-card\" data-anchor=\"{i['anchor']}\" "
            f"style=\"border-left:4px solid {color}\">"
            f"<div class=\"term\"><strong>{html.escape(i['issue'])}</strong></div>"
            f"<div><em>Replacement:</em> {html.escape(i['replacement'] or '—')}</div>"
            f"<div>{html.escape(i['explanation'])}</div>"
            f"<div class=\"loc\">{html.escape(i['location'])}</div>"
            f"</a>"
        )
    right_html = "".join(right_items) or "<div class='no-issues'>No issues found 🎉</div>"

    PAGE_TMPL = Template("""
    <style>
      html { scroll-behavior:smooth; }
      .wrap { display:grid; grid-template-columns:1fr 420px; gap:20px; }
      .doc, .issues {
        height:650px; overflow:auto; border:1px solid #e6e6e6;
        padding:16px; border-radius:8px; background:#fff;
      }
      .issues { background:#fafafa; }
      .issue-card {
        display:block; text-decoration:none; color:#000;
        background:#fff; border:1px solid #eee;
        border-radius:6px; padding:10px; margin-bottom:10px;
      }
      .doc-table td { border:1px solid #eee; padding:8px 10px; }
      @keyframes flashBg {
        0% { background:#ffe8a0; }
        100% { background:transparent; }
      }
      .flash { animation: flashBg 1.2s ease-out; }
    </style>

    <div class="wrap">
      <div class="doc" id="doc-pane">
        $left_html
      </div>
      <div class="issues">
        <h4>Issues</h4>
        $right_html
      </div>
    </div>

    <script>
    (function(){
      const docPane = document.getElementById("doc-pane");
      function offsetTop(el, ancestor){
        let top = 0;
        while (el && el !== ancestor){
          top += el.offsetTop;
          el = el.offsetParent;
        }
        return top;
      }
      function flash(el){
        el.classList.remove("flash");
        void el.offsetWidth;
        el.classList.add("flash");
      }
      document.querySelectorAll(".issue-card").forEach(card => {
        card.addEventListener("click", e => {
          e.preventDefault();
          const id = card.dataset.anchor;
          const target = document.getElementById(id);
          if (!target || !docPane) return;
          docPane.scrollTo({ top: offsetTop(target, docPane) - 12, behavior:"smooth" });
          flash(target);
        });
      });
    })();
    </script>
    """)
    full_html = PAGE_TMPL.substitute(left_html=left_html, right_html=right_html)
    return full_html, issues_sorted


# ===========================
# NEW: Analyze raw pasted text
# ===========================
def analyze_text_input(text_input: str):
    """
    Analyze raw text pasted/typed by the user and return (full_html, issues_sorted)
    using the same rendering pipeline as the .docx analyzer.
    Blank lines split paragraphs.
    """
    raw = text_input or ""
    blocks = [blk for blk in re.split(r"\n\s*\n", raw.strip())] if raw.strip() else []

    rules = flatten_rules()
    left_parts, issues = [], []

    for i, paragraph in enumerate(blocks, 1):
        text = _normalize_for_match(paragraph)
        loc  = f"Paragraph {i}"

        rule_matches = find_matches(text, rules, loc, f"tp{i}")
        dict_issues  = find_non_dictionary_words(text, loc, f"tp{i}")
        ukus_issues  = find_non_us_words(text, loc, f"tp{i}")
        caps_sent    = find_all_caps_sentences(text, loc, f"tp{i}")
        caps_words   = find_all_caps_words(text, loc, f"tp{i}")

        # Match DOCX logic for consistency.
        caps_words = _suppress_cap_words_inside_cap_sent(caps_sent, caps_words)
        if caps_sent or _is_block_mostly_all_caps(text):
            caps_words = []

        combined = _dedupe_by_span(rule_matches + ukus_issues + dict_issues + caps_sent + caps_words)
        issues.extend(combined)
        left_parts.append(f"<p>{render_text(text, combined) or '&nbsp;'}</p>")

    left_html = "".join(left_parts) or "<p>&nbsp;</p>"

    # Right pane
    issues_sorted = sorted(issues, key=_severity_key)
    right_items = []
    for i in issues_sorted:
        color = "#ff4d4d" if i["category"] == "style_guide_rule" else "#ffcc00"
        right_items.append(
            f"<a href=\"#{i['anchor']}\" class=\"issue-card\" data-anchor=\"{i['anchor']}\" "
            f"style=\"border-left:4px solid {color}\">"
            f"<div class=\"term\"><strong>{html.escape(i['issue'])}</strong></div>"
            f"<div><em>Replacement:</em> {html.escape(i['replacement'] or '—')}</div>"
            f"<div>{html.escape(i['explanation'])}</div>"
            f"<div class=\"loc\">{html.escape(i['location'])}</div>"
            f"</a>"
        )
    right_html = "".join(right_items) or "<div class='no-issues'>No issues found 🎉</div>"

    PAGE_TMPL = Template("""
    <style>
      html { scroll-behavior:smooth; }
      .wrap { display:grid; grid-template-columns:1fr 420px; gap:20px; }
      .doc, .issues {
        height:650px; overflow:auto; border:1px solid #e6e6e6;
        padding:16px; border-radius:8px; background:#fff;
      }
      .issues { background:#fafafa; }
      .issue-card {
        display:block; text-decoration:none; color:#000;
        background:#fff; border:1px solid #eee;
        border-radius:6px; padding:10px; margin-bottom:10px;
      }
      .doc-table td { border:1px solid #eee; padding:8px 10px; }
      @keyframes flashBg {
        0% { background:#ffe8a0; }
        100% { background:transparent; }
      }
      .flash { animation: flashBg 1.2s ease-out; }
    </style>

    <div class="wrap">
      <div class="doc" id="doc-pane">
        $left_html
      </div>
      <div class="issues">
        <h4>Issues</h4>
        $right_html
      </div>
    </div>

    <script>
    (function(){
      const docPane = document.getElementById("doc-pane");
      function offsetTop(el, ancestor){
        let top = 0;
        while (el && el !== ancestor){
          top += el.offsetTop;
          el = el.offsetParent;
        }
        return top;
      }
      function flash(el){
        el.classList.remove("flash");
        void el.offsetWidth;
        el.classList.add("flash");
      }
      document.querySelectorAll(".issue-card").forEach(card => {
        card.addEventListener("click", e => {
          e.preventDefault();
          const id = card.dataset.anchor;
          const target = document.getElementById(id);
          if (!target || !docPane) return;
          docPane.scrollTo({ top: offsetTop(target, docPane) - 12, behavior:"smooth" });
          flash(target);
        });
      });
    })();
    </script>
    """)
    full_html = PAGE_TMPL.substitute(left_html=left_html, right_html=right_html)
    return full_html, issues_sorted

# ===========================
# TABS
# ===========================
tab_check, tab_rules = st.tabs(["📄 Style Checker", "📋 Add/Edit Rules"])

with tab_check:
    st.markdown("Use either **Paste text** or **Upload .docx** to analyze content.")
    sub_paste, sub_docx = st.tabs(["✍️ Paste text", "📤 Upload .docx"])

    with sub_paste:
        default_placeholder = (
            "Paste or type your text here.\n\n"
            "Example:\n"
            "THIS IS AN ALL CAPS SENTENCE FOR TESTING.\n\n"
            "Our favourite neighbour visited the centre yesterday. "
            "We also noticed an envionmentt issue."
        )
        text_input = st.text_area(
            "Text to analyze",
            value=default_placeholder,
            height=220,
            help="Blank lines split paragraphs. The checker uses the same rules as the .docx path."
        )
        c1, c2 = st.columns([1, 1])
        run_paste = c1.button("Analyze text", type="primary", use_container_width=True)
        if c2.button("Clear", type="secondary", use_container_width=True):
            st.session_state.pop("last_text_results", None)
            st.rerun()

        if run_paste and text_input.strip():
            page_html, issues = analyze_text_input(text_input)
            st.session_state["last_text_results"] = (page_html, issues)

        if "last_text_results" in st.session_state:
            page_html, issues = st.session_state["last_text_results"]
            components.html(page_html, height=720, scrolling=True)

            st.subheader("🧾 Summary Table (pasted text)")
            st.dataframe(
                [
                    {
                        "issue": i["issue"],
                        "replacement": i["replacement"] or "—",
                        "explanation": i["explanation"],
                        "location": i["location"],
                    }
                    for i in issues
                ],
                use_container_width=True,
            )

    with sub_docx:
        uploaded = st.file_uploader("Upload Word document (.docx)", type=["docx"])
        if uploaded:
            page_html, issues = analyze_inline(uploaded.read())
            components.html(page_html, height=720, scrolling=True)

            st.subheader("🧾 Summary Table (.docx)")
            st.dataframe(
                [
                    {
                        "issue": i["issue"],
                        "replacement": i["replacement"] or "—",
                        "explanation": i["explanation"],
                        "location": i["location"],
                    }
                    for i in issues
                ],
                use_container_width=True,
            )

with tab_rules:
    def _reload_rules():
        try:
            load_rules_sheets.clear()
        except Exception:
            pass
        st.session_state.rules = load_rules()
    st.button(
        "↻ Reload rules from source",
        help="Clear cache and reload from Google Sheets (or local JSON).",
        on_click=_reload_rules,
        type="secondary",
    )

    ensure_state()

    if SHEETS_ENABLED:
        st.success("Using Google Sheets as the source of truth for rules.")
    else:
        st.info("Using local JSON files (Google Sheets not configured).")

    st.subheader("➕ Add New Rule")
    with st.form("add_rule"):
        category = st.selectbox(
            "Category", ["style_guide_rule", "style_guide_caution"],
            format_func=lambda x: x.replace("_", " ").title(),
        )
        match = st.text_input("Word or phrase to flag")
        replacement = st.text_input("Suggested replacement (optional)")
        message = st.text_input("Explanation shown to users")
        case_sensitive = st.checkbox("Case sensitive match", value=False)
        submitted = st.form_submit_button("Add rule")

        if submitted and match and message:
            rules_state = st.session_state.get("rules", {"style_guide_rule": [], "style_guide_caution": []})
            rules_state.setdefault(category, []).insert(
                0,
                {"match": match, "replace_with": replacement or None,
                 "message": message, "case_sensitive": case_sensitive}
            )
            st.session_state.rules = rules_state
            save_rules(st.session_state.rules)
            st.success("Rule added successfully.")
            st.rerun()

    st.divider()
    st.subheader("📋 Existing Rules")
    for cat in ("style_guide_rule", "style_guide_caution"):
        st.markdown(f"### {cat.replace('_', ' ').title()}")
        rules_state = st.session_state.get("rules", {"style_guide_rule": [], "style_guide_caution": []})
        rules_list = rules_state.get(cat, [])
        if not rules_list:
            st.info("No rules in this category yet.")
            continue
        for idx, rule in enumerate(rules_list):
            cols = st.columns([5, 2, 1])
            with cols[0]:
                if st.session_state.edit_rule == (cat, idx):
                    new_match = st.text_input("Match", rule["match"], key=f"edit_match_{cat}_{idx}")
                    new_repl  = st.text_input("Replacement", rule.get("replace_with") or "", key=f"edit_repl_{cat}_{idx}")
                    new_msg   = st.text_input("Message", rule["message"], key=f"edit_msg_{cat}_{idx}")
                    new_cs    = st.checkbox("Case sensitive", value=_coerce_bool(rule.get("case_sensitive", False)),
                                            key=f"edit_cs_{cat}_{idx}")
                else:
                    st.markdown(
                        f"**Match:** `{rule['match']}`  \n"
                        f"**Replacement:** {rule.get('replace_with') or '—'}  \n"
                        f"**Message:** {rule['message']}  \n"
                        f"**Case sensitive:** {_coerce_bool(rule.get('case_sensitive', False))}"
                    )
            with cols[1]:
                if st.session_state.edit_rule == (cat, idx):
                    if st.button("💾 Save", key=f"save_{cat}_{idx}"):
                        rule["match"] = new_match
                        rule["replace_with"] = new_repl or None
                        rule["message"] = new_msg
                        rule["case_sensitive"] = bool(new_cs)
                        st.session_state.rules = rules_state
                        save_rules(st.session_state.rules)
                        st.session_state.edit_rule = None
                        st.rerun()
                    if st.button("✖ Cancel", key=f"cancel_{cat}_{idx}"):
                        st.session_state.edit_rule = None
                        st.rerun()
                else:
                    if st.button("✏ Edit", key=f"edit_{cat}_{idx}"):
                        st.session_state.edit_rule = (cat, idx)
                        st.rerun()
            with cols[2]:
                if st.button("🗑 Delete", key=f"delete_{cat}_{idx}"):
                    rules_state[cat].pop(idx)
                    st.session_state.rules = rules_state
                    save_rules(st.session_state.rules)
                    st.rerun()

    # ---------------- Acronym Whitelist (ALL‑CAPS) ----------------
    st.divider()
    st.subheader("🔠 Acronym Whitelist (ALL‑CAPS)")
    st.caption(
        "Words in ALL‑CAPS are normally flagged as a *caution*. "
        "Add acronyms here (letters only) to **allow** them without warnings. "
        f"This list is stored in the Google Sheets tab `{CAPS_WS_NAME}` when Sheets is enabled, "
        "or in `Rules/caps_whitelist.json` locally."
    )

    with st.form("add_caps_whitelist"):
        new_acro = st.text_input("Add acronym (letters only, e.g., 'UN', 'SDG')").strip().upper()
        add_it = st.form_submit_button("Add acronym")
        if add_it:
            if new_acro and new_acro.isalpha() and len(new_acro) >= 2:
                ac = set(st.session_state.get("caps_whitelist", set()))
                ac.add(new_acro)
                st.session_state.caps_whitelist = ac
                save_caps_whitelist(ac)
                st.success(f"Added '{new_acro}' to the acronym whitelist.")
                st.rerun()
            else:
                st.error("Please enter letters only (A‑Z), at least 2 characters.")

    acros = sorted(st.session_state.get("caps_whitelist", set()))
    if not acros:
        st.info("No custom acronyms yet.")
    else:
        cols = st.columns(6)
        for i, acr in enumerate(acros):
            col = cols[i % 6]
            with col:
                st.write(f"**{acr}**")
                if st.button("Remove", key=f"rm_acr_{acr}"):
                    ac = set(st.session_state.get("caps_whitelist", set()))
                    if acr in ac:
                        ac.remove(acr)
                        st.session_state.caps_whitelist = ac
                        save_caps_whitelist(ac)
                        st.toast(f"Removed '{acr}'")
                        st.rerun()

    def _reload_caps_whitelist():
        try:
            load_caps_whitelist_sheets.clear()
        except Exception:
            pass
        st.session_state.caps_whitelist = load_caps_whitelist()

    st.button("↻ Reload acronym whitelist", on_click=_reload_caps_whitelist, type="secondary")
