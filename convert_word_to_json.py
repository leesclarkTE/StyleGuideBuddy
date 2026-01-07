from docx import Document
import json

# Path to your Word style guide
doc_path = "Textile_Exchange_Style_Guide.docx"

doc = Document(doc_path)
rules = {}

# Helper function to process a line
def process_line(line):
    line = line.strip()
    if not line or "→" not in line:
        return
    term, right = line.split("→", 1)
    term = term.strip()
    right = right.strip()
    if right.lower().startswith("message:"):
        rules[term] = {"message": right.replace("message:", "").strip(), "auto_fix": False}
    else:
        rules[term] = {"replacement": right, "auto_fix": True}

# 1️⃣ Read paragraphs
for para in doc.paragraphs:
    for line in para.text.split("\n"):
        process_line(line)

# 2️⃣ Read tables (if rules are in a table)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for line in cell.text.split("\n"):
                process_line(line)

# Save JSON
with open("rules/terminology.json", "w") as f:
    json.dump(rules, f, indent=2)

print(f"terminology.json created with {len(rules)} rules!")

